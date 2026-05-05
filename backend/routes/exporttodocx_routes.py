from flask import Blueprint, jsonify, send_file
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from services import ServiceFactory

exportdocx_bp = Blueprint('exportdocx', __name__)

# ===== HELPERS =====
def set_cell_color(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def calculate_year(studyplan_year, semester_number, term):
    return studyplan_year + (semester_number - 1) // 2 + (1 if term == "V" else 0)


def generate_studyplan_docx(studyprogram, studyplans, course_to_package):
    doc = Document()

    doc.add_heading(f"{studyprogram.name}", level=1)
    doc.add_paragraph(f"Institution: {studyprogram.institute.name}")
    doc.add_paragraph(f"Ansvarlig: {studyprogram.institute.ansvarlig}")
    doc.add_paragraph(f"Degree Type: {studyprogram.degree_type}")
    doc.add_paragraph("\n")

    first_studyplan = studyplans[0]
    start_year = first_studyplan['year']

    end_year = calculate_year(
        start_year,
        first_studyplan['semesters'][-1]['semester_number'],
        first_studyplan['semesters'][-1]['term']
    )

    doc.add_heading("Studieplanmatrise", level=2)
    doc.add_paragraph(f"Studieplan {start_year} - {end_year} for {studyprogram.name}")

    # ===== FARGER =====
    color_palette = [
        "FFFF00", "FFA500", "00B050",
        "0070C0", "FF0000", "FF66CC", "FFFFFF", "00FFFF", "C00000",
    ]

    package_colors = {}
    color_index = 0

    for pkg in course_to_package.values():
        name = pkg["name"]
        if name not in package_colors:
            package_colors[name] = color_palette[color_index % len(color_palette)]
            color_index += 1

    package_colors["Andre emner"] = "D9D9D9"
    package_colors["VALGEMNE"] = "0070C0"

    # ===== TABELL =====
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Semester"
    header_cells[1].text = "10 sp"
    header_cells[2].text = "10 sp"
    header_cells[3].text = "10 sp"

    # ===== MATRISE =====
    for semester in first_studyplan['semesters']:
        semester_year = calculate_year(
            start_year,
            semester['semester_number'],
            semester['term']
        )

        blocks = [[], [], []]
        current_block = 0
        current_block_credits = 0

        for course in semester['semester_courses']:
            if course['is_elective']:
                blocks = [[{
                    "text": "VALGEMNE",
                    "credits": 0,
                    "pkg": "VALGEMNE"
                }] for _ in range(3)]
                break

            course_id = course.get("id")
            pkg = course_to_package.get(course_id)
            pkg_name = pkg["name"] if pkg else "Andre emner"

            course_data = {
                "text": f"{course['courseCode']} ({course['credits']} sp)",
                "credits": course['credits'] or 0,
                "pkg": pkg_name
            }

            if current_block_credits + course_data["credits"] <= 10:
                blocks[current_block].append(course_data)
                current_block_credits += course_data["credits"]
            else:
                current_block += 1
                if current_block < 3:
                    blocks[current_block].append(course_data)
                    current_block_credits = course_data["credits"]

        max_rows = max(len(b) for b in blocks)
        semester_cells = []

        for i in range(max_rows):
            row = table.add_row()
            row_cells = row.cells

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "500")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            if i == 0:
                row_cells[0].text = f"{semester['semester_number']} ({semester['term']}-{semester_year})"
            semester_cells.append(row_cells[0])

            for col_idx in range(3):
                cell = row_cells[col_idx + 1]

                if i < len(blocks[col_idx]):
                    c = blocks[col_idx][i]
                    cell.text = c["text"]

                    color = package_colors.get(c["pkg"], "FFFFFF")
                    set_cell_color(cell, color)

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)

                    tcPr = cell._tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    for m in ["top", "left", "bottom", "right"]:
                        node = OxmlElement(f"w:{m}")
                        node.set(qn('w:w'), "100")
                        node.set(qn('w:type'), 'dxa')
                        tcMar.append(node)
                    tcPr.append(tcMar)

        for i in range(1, len(semester_cells)):
            semester_cells[0].merge(semester_cells[i])

    doc.add_paragraph("\n")

    # ===== LEGENDE =====
    legend_table = doc.add_table(rows=1, cols=len(package_colors))

    for i, (pkg, color) in enumerate(package_colors.items()):
        cell = legend_table.rows[0].cells[i]
        cell.text = pkg
        set_cell_color(cell, color)

    doc.add_paragraph("\n")

    # ===== EMNEOVERSIKT =====
    doc.add_heading("Emneoversikt", level=2)

    for semester in first_studyplan['semesters']:
        semester_year = calculate_year(
            start_year,
            semester['semester_number'],
            semester['term']
        )

        doc.add_heading(
            f"Semester {semester['semester_number']} ({semester['term']} {semester_year})",
            level=2
        )

        courses_by_package = {}

        for course in semester['semester_courses']:
            course_id = course.get('id')

            if course['is_elective']:
                category = course.get('category')

                if category:
                    pkg_name = f"{category.get('name')}"
                else:
                    pkg_name = "Valgfag (uten gruppe)"
            else:
                pkg = course_to_package.get(course_id)
                pkg_name = pkg["name"] if pkg else "Andre emner"

            courses_by_package.setdefault(pkg_name, []).append(course)

        sorted_packages = sorted(
            courses_by_package.items(),
            key=lambda x: (x[0] != "Andre emner", x[0])
        )

        for pkg_name, courses_in_pkg in sorted_packages:
            total_credits = sum(course.get("credits") or 0 for course in courses_in_pkg)

            pkg_para = doc.add_paragraph(f"{pkg_name} (sp: {total_credits})")
            pkg_para.style = "Heading 3"
            pkg_para.paragraph_format.left_indent = Inches(0.3)

            for course in courses_in_pkg:
                course_para = doc.add_paragraph(
                    f"{course['courseCode']} - {course['name']} ({course['credits']} sp)"
                )
                course_para.paragraph_format.left_indent = Inches(0.6)

    return doc


# ===== ROUTE =====
@exportdocx_bp.route('/<int:studyprogram_id>', methods=['GET'])
def export_to_docx(studyprogram_id):
    try:
        studyplan_service = ServiceFactory.get_studyplan_service()
        studyprogram_service = ServiceFactory.get_studyprogram_service()
        coursepackage_service = ServiceFactory.get_coursepackage_service()

        studyprogram = studyprogram_service.get_studyprogram_by_id(studyprogram_id)

        if not studyprogram:
            return jsonify({"error": "Study program not found"}), 404

        studyplans = studyplan_service.get_plans_for_export(studyprogram_id)

        if not studyplans:
            return jsonify({"error": "No studyplans found"}), 400

        selected_plan = studyplans[0]
        selected_plan_id = selected_plan.get('id')

        if not selected_plan_id:
            packages = []
        else:
            packages = coursepackage_service.get_course_packages_by_studyplan(selected_plan_id)

        course_to_package = {}

        for pkg in packages:
            courses_in_pkg = coursepackage_service.get_courses_in_package(pkg.id)

            for course in courses_in_pkg:
                course_to_package[course.id] = {
                    "name": pkg.name,
                    "type": pkg.packagetype
                }

        doc = generate_studyplan_docx(
            studyprogram,
            studyplans,
            course_to_package
        )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"studyplan_{studyprogram_id}_{studyplans[0]['year']}.docx"

        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500